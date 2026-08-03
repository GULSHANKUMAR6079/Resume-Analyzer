import io
import sys
import logging
from typing import Tuple, Optional
import pdfplumber
import PyPDF2
from docx import Document

from backend.core.config import (
    MAX_FILE_SIZE_BYTES,
    MAX_FILE_SIZE_MB,
    SUPPORTED_MIME_TYPES,
)

logger = logging.getLogger('ats_resume_scorer')

class FileParsingError(Exception):
    pass

class FileValidationError(Exception):
    pass

def validate_file(file_data: bytes, filename: str) -> Tuple[bool, str, Optional[str]]:
    file_size_bytes = len(file_data)
    if file_size_bytes > MAX_FILE_SIZE_BYTES:
        size_mb = file_size_bytes / (1024 * 1024)
        return False, (
            f"File size ({size_mb:.2f} MB) exceeds maximum allowed size of {MAX_FILE_SIZE_MB} MB."
        ), None

    if file_size_bytes == 0:
        return False, "Uploaded file is empty. Please upload a valid document.", None

    # Detect MIME type safely across platforms
    mime_type = None
    try:
        if sys.platform == 'win32':
            try:
                import magic
                mime_type = magic.from_buffer(file_data, mime=True)
            except Exception:
                pass
        else:
            import magic
            mime_type = magic.from_buffer(file_data, mime=True)
    except Exception:
        pass

    # Fallback MIME detection by file extension if python-magic is unavailable or fails
    if not mime_type:
        ext = filename.lower().split('.')[-1] if '.' in filename else ''
        if ext == 'pdf':
            mime_type = 'application/pdf'
        elif ext == 'docx':
            mime_type = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        elif ext == 'doc':
            mime_type = 'application/msword'

    if mime_type not in SUPPORTED_MIME_TYPES:
        supported = ', '.join(SUPPORTED_MIME_TYPES.keys()).upper()
        return False, f"Unsupported file type: {mime_type}. Supported types are: {supported}", None

    return True, "", SUPPORTED_MIME_TYPES[mime_type]

def _extract_pdf_hyperlinks(file_data: bytes) -> str:
    urls = []
    try:
        reader = PyPDF2.PdfReader(io.BytesIO(file_data))
        for page in reader.pages:
            if '/Annots' not in page:
                continue
            for annot_ref in page['/Annots']:
                try:
                    annot = annot_ref.get_object()
                    if annot.get('/Subtype') != '/Link':
                        continue
                    action = annot.get('/A', {})
                    uri = action.get('/URI', '')
                    if uri:
                        if isinstance(uri, bytes):
                            uri = uri.decode('utf-8', errors='ignore')
                        uri = uri.strip()
                        if uri.startswith('http'):
                            urls.append(uri)
                except Exception:
                    pass
    except Exception:
        pass
    return '\n'.join(list(set(urls)))

def _extract_pdf_with_pdfplumber(file_data: bytes) -> str:
    text_parts = []
    with pdfplumber.open(io.BytesIO(file_data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    text = '\n'.join(text_parts).strip()
    if not text:
        raise FileParsingError("pdfplumber extracted no text from the PDF.")

    hyperlinks = _extract_pdf_hyperlinks(file_data)
    if hyperlinks:
        text += '\n' + hyperlinks
    return text

def _extract_pdf_with_pypdf2(file_data: bytes) -> str:
    text_parts = []
    pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    text = '\n'.join(text_parts).strip()
    if not text:
        raise FileParsingError("PyPDF2 extracted no text from the PDF.")

    hyperlinks = _extract_pdf_hyperlinks(file_data)
    if hyperlinks:
        text += '\n' + hyperlinks
    return text

def extract_text_from_pdf(file_data: bytes) -> str:
    try:
        return _extract_pdf_with_pdfplumber(file_data)
    except Exception as exc:
        logger.warning(f"pdfplumber failed: {exc}. Trying PyPDF2 fallback...")
        try:
            return _extract_pdf_with_pypdf2(file_data)
        except Exception as exc2:
            logger.error(f"PyPDF2 fallback also failed: {exc2}")
            raise FileParsingError(
                "Could not extract selectable text from PDF. "
                "The PDF may be corrupted, password protected, or contain scanned images."
            )

def extract_text_from_docx(file_data: bytes) -> str:
    try:
        doc = Document(io.BytesIO(file_data))
        text_parts = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)

        text = '\n'.join(text_parts).strip()
        if not text:
            raise FileParsingError("No text found in DOCX file.")

        # Extract hyperlinks from DOCX XML relations
        try:
            for rel in doc.part.rels.values():
                if 'hyperlink' in rel.reltype.lower():
                    url = rel._target
                    if isinstance(url, str) and url.startswith('http'):
                        text += '\n' + url
        except Exception:
            pass

        return text
    except FileParsingError:
        raise
    except Exception as e:
        raise FileParsingError(f"Failed to parse DOCX file: {e}")

def parse_resume_file(file_data: bytes, filename: str) -> Tuple[str, dict]:
    is_valid, error_msg, file_type = validate_file(file_data, filename)
    if not is_valid:
        raise FileValidationError(error_msg)

    if file_type == 'pdf':
        text = extract_text_from_pdf(file_data)
    elif file_type == 'docx':
        text = extract_text_from_docx(file_data)
    elif file_type == 'doc':
        raise FileParsingError("Legacy .doc format is not supported. Please convert to .docx or .pdf.")
    else:
        raise FileValidationError(f"Unsupported file format: {file_type}")

    metadata = {
        'filename': filename,
        'file_type': file_type,
        'file_size_bytes': len(file_data),
        'text_length': len(text),
        'success': True,
    }
    return text, metadata
